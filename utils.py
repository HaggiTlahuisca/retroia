"""Funciones utilitarias para manejo de archivos, limpieza y Moodle."""

import io
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def now_slug() -> str:
    """Devuelve la fecha y hora actual en formato compacto (ej. 20260821_153022)."""
    return datetime.now().strftime("%Y%m%d_%H%M%S")

def sanitize_filename(name: str) -> str:
    """Elimina caracteres no válidos para nombres de archivo."""
    return re.sub(r'[^\w\s-]', '', name).strip().replace(' ', '_')

def get_activity_code(name: str) -> str:
    """Genera un código corto para el nombre del archivo (ej. AI1, PI, FI)."""
    lower_name = name.lower()
    
    if "proyecto integrador" in lower_name:
        return "PI"
    if "foro de integración" in lower_name:
        return "FI"
    if "actividad integradora" in lower_name:
        # Diccionario para traducir letras a números
        numeros = {"uno": 1, "dos": 2, "tres": 3, "cuatro": 4, "cinco": 5, "seis": 6}
        
        # 1. Buscar si hay un dígito numérico explícito (ej. "Actividad integradora 1")
        match = re.search(r'\d+', lower_name)
        if match:
            return f"AI{match.group()}"
            
        # 2. Buscar si el número está escrito con letras (ej. "Actividad integradora uno")
        for palabra, num in numeros.items():
            if palabra in lower_name:
                return f"AI{num}"
                
        return "AI"
        
    return "Gen"

def docx_bytes(student_name: str, text: str) -> bytes:
    """Genera un archivo Word en memoria y retorna sus bytes."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    lineas = text.split("\n")
    firmas = ["haggi de jesús tlahuisca hernández", "asesor virtual", "21d28277", "cordialmente.", "con afecto.", "atentamente."]

    for linea in lineas:
        limpia = linea.strip()
        if not limpia:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        lower = limpia.lower()
        es_grupo = bool(re.match(r"^m\d{1,2}c\d{1,2}g\d{1,3}-\d{3}$", lower))

        if lower in firmas or es_grupo or lower.startswith("criterio ") or lower.startswith("apreciable"):
            texto_plano = limpia.replace("**", "").replace("##", "")
            run = p.add_run(texto_plano)
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.font.bold = True
            continue

        tokens = re.split(r"(\*\*.*?\*\*)", limpia)
        for token in tokens:
            if not token:
                continue
            if token.startswith("**") and token.endswith("**"):
                run = p.add_run(token[2:-2])
                run.font.name = "Arial"
                run.font.size = Pt(12)
                run.font.bold = True
            else:
                run = p.add_run(token)
                run.font.name = "Arial"
                run.font.size = Pt(12)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()

def feedback_to_moodle_html(text: str) -> str:
    """Convierte el texto plano en HTML simple para Moodle."""
    html = ""
    for parrafo in text.split("\n\n"):
        if parrafo.strip():
            p = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', parrafo.strip())
            p = p.replace("\n", "<br>")
            html += f"<p>{p}</p>\n"
    return html
